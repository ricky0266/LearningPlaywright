from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ── Styles helper functions ──────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_normal(doc, text):
    doc.add_paragraph(text)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x88)

def add_spacer(doc):
    doc.add_paragraph("")

# ── Title ────────────────────────────────────────────────────────────────────
add_title(doc, "Playwright CLI - Complete Command Reference")
add_normal(doc, f"Version: 1.58.2  |  Generated: {datetime.date.today().strftime('%d-%m-%Y')}")
add_spacer(doc)

# ── 1. open ──────────────────────────────────────────────────────────────────
add_heading(doc, "1. open - Open a page in browser")
add_code(doc, "npx playwright open [options] [url]")
add_normal(doc, "Options:")
add_code(doc, "  -b, --browser <browser>        chromium | firefox | webkit  (default: chromium)")
add_code(doc, "  --channel <channel>            Browser channel e.g. chrome, msedge")
add_code(doc, "  --color-scheme <scheme>        light | dark | no-preference")
add_code(doc, "  --device <deviceName>          Emulate device e.g. 'iPhone 13'")
add_code(doc, "  --geolocation <coordinates>    'lat,lon'")
add_code(doc, "  --ignore-https-errors          Ignore HTTPS errors")
add_code(doc, "  --lang <language>              Language e.g. en-US")
add_code(doc, "  --load-storage <filename>      Load storage state from file")
add_code(doc, "  --proxy-server <proxy>         Proxy server e.g. http://myproxy:3128")
add_code(doc, "  --save-storage <filename>      Save storage state after session")
add_code(doc, "  --save-trace <filename>        Save trace file")
add_code(doc, "  --timezone <timezone>          Time zone e.g. 'America/New_York'")
add_code(doc, "  --timeout <timeout>            Timeout in ms")
add_code(doc, "  --user-agent <ua string>       Custom user agent")
add_code(doc, "  --viewport-size <size>         e.g. '1280, 720'")
add_spacer(doc)

# ── 2. codegen ───────────────────────────────────────────────────────────────
add_heading(doc, "2. codegen - Record actions and generate code")
add_code(doc, "npx playwright codegen [options] [url]")
add_normal(doc, "Options:")
add_code(doc, "  -o, --output <file>            Save script to file e.g. tests/recorded.spec.js")
add_code(doc, "  --target <language>            javascript | typescript | python | python-async | java | csharp")
add_code(doc, "  -b, --browser <browser>        chromium | firefox | webkit")
add_code(doc, "  --channel <channel>            Browser channel e.g. chrome, msedge")
add_code(doc, "  --color-scheme <scheme>        light | dark | no-preference")
add_code(doc, "  --device <deviceName>          Emulate device e.g. 'Pixel 5'")
add_code(doc, "  --geolocation <coordinates>    'lat,lon'")
add_code(doc, "  --ignore-https-errors          Ignore HTTPS errors")
add_code(doc, "  --lang <language>              Language e.g. en-US")
add_code(doc, "  --load-storage <filename>      Load auth/cookies from file")
add_code(doc, "  --save-storage <filename>      Save auth/cookies to file")
add_code(doc, "  --save-trace <filename>        Save trace file")
add_code(doc, "  --timezone <timezone>          Time zone")
add_code(doc, "  --viewport-size <size>         '1280, 720'")
add_code(doc, "  --user-agent <ua>              Custom user agent")
add_code(doc, "  --proxy-server <proxy>         Proxy server")
add_spacer(doc)

# ── 3. install ───────────────────────────────────────────────────────────────
add_heading(doc, "3. install - Install browsers")
add_code(doc, "npx playwright install [options] [browser...]")
add_normal(doc, "Browsers: chromium | firefox | webkit | chromium-headless-shell")
add_normal(doc, "Options:")
add_code(doc, "  --with-deps     Install OS dependencies for browsers")
add_code(doc, "  --dry-run       Just list what would be installed")
add_code(doc, "  --force         Force reinstall even if already installed")
add_spacer(doc)

# ── 4. cr / ff / wk ──────────────────────────────────────────────────────────
add_heading(doc, "4. cr / ff / wk - Open in specific browser")
add_code(doc, "npx playwright cr [options] [url]    # Chromium")
add_code(doc, "npx playwright ff [options] [url]    # Firefox")
add_code(doc, "npx playwright wk [options] [url]    # WebKit")
add_normal(doc, "(Same options as the 'open' command)")
add_spacer(doc)

# ── 5. screenshot ─────────────────────────────────────────────────────────────
add_heading(doc, "5. screenshot - Capture page screenshot")
add_code(doc, "npx playwright screenshot [options] <url> <filename>")
add_normal(doc, "Options:")
add_code(doc, "  -b, --browser <browser>        chromium | firefox | webkit")
add_code(doc, "  --device <deviceName>          Emulate device")
add_code(doc, "  --color-scheme <scheme>        light | dark")
add_code(doc, "  --full-page                    Capture full scrollable page")
add_code(doc, "  --wait-for-selector <selector> Wait for element before screenshot")
add_code(doc, "  --viewport-size <size>         '1280, 720'")
add_code(doc, "  --geolocation <coordinates>    'lat,lon'")
add_code(doc, "  --lang <language>              Language")
add_code(doc, "  --proxy-server <proxy>         Proxy")
add_code(doc, "  --ignore-https-errors          Ignore HTTPS errors")
add_code(doc, "  --save-trace <filename>        Save trace")
add_code(doc, "  --timeout <ms>                 Timeout")
add_spacer(doc)

# ── 6. pdf ────────────────────────────────────────────────────────────────────
add_heading(doc, "6. pdf - Save page as PDF")
add_code(doc, "npx playwright pdf [options] <url> <filename>")
add_normal(doc, "Options:")
add_code(doc, "  --wait-for-selector <selector> Wait for element")
add_code(doc, "  --viewport-size <size>         Viewport size")
add_code(doc, "  --timeout <ms>                 Timeout")
add_code(doc, "  -b, --browser <browser>        Only chromium supports PDF")
add_spacer(doc)

# ── 7. test ───────────────────────────────────────────────────────────────────
add_heading(doc, "7. test - Run tests")
add_code(doc, "npx playwright test [options] [test-filter...]")
add_normal(doc, "Options:")
add_code(doc, "  --headed                       Run in headed mode (show browser)")
add_code(doc, "  --ui                           Open UI mode (visual runner)")
add_code(doc, "  --debug                        Run with Playwright Inspector")
add_code(doc, "  -g, --grep <pattern>           Only run tests matching pattern")
add_code(doc, "  --grep-invert <pattern>        Skip tests matching pattern")
add_code(doc, "  -x                             Stop after first failure")
add_code(doc, "  --retries <number>             Retry failed tests N times")
add_code(doc, "  --workers <number>             Parallel workers count")
add_code(doc, "  --timeout <ms>                 Test timeout")
add_code(doc, "  --reporter <reporter>          list | dot | html | json | line")
add_code(doc, "  --project <name>               Run specific project from config")
add_code(doc, "  --config <file>                Path to config file")
add_code(doc, "  --ignore-snapshots             Ignore screenshot diffs")
add_code(doc, "  --update-snapshots             Update snapshot files")
add_code(doc, "  --pass-with-no-tests           Exit 0 if no tests found")
add_code(doc, "  --trace <mode>                 on | off | retain-on-failure | on-first-retry")
add_code(doc, "  --video <mode>                 on | off | retain-on-failure")
add_code(doc, "  --screenshot <mode>            on | off | only-on-failure")
add_code(doc, "  --forbid-only                  Fail if test.only is used")
add_code(doc, "  --fully-parallel               Run all tests in parallel")
add_code(doc, "  --last-failed                  Re-run only last failed tests")
add_code(doc, "  --shard <current/all>          Shard tests e.g. 1/3")
add_spacer(doc)

# ── 8. show-report ────────────────────────────────────────────────────────────
add_heading(doc, "8. show-report - View HTML test report")
add_code(doc, "npx playwright show-report [options] [report-folder]")
add_normal(doc, "Options:")
add_code(doc, "  --host <host>        Host to serve report  (default: localhost)")
add_code(doc, "  --port <port>        Port to serve report  (default: 9323)")
add_spacer(doc)

# ── 9. show-trace ─────────────────────────────────────────────────────────────
add_heading(doc, "9. show-trace - View trace file")
add_code(doc, "npx playwright show-trace [options] [trace-file.zip]")
add_normal(doc, "Options:")
add_code(doc, "  --browser <browser>   Browser to use for viewer")
add_code(doc, "  --port <port>         Port for trace viewer server")
add_code(doc, "  --host <host>         Host for trace viewer server")
add_spacer(doc)

# ── 10. merge-reports ─────────────────────────────────────────────────────────
add_heading(doc, "10. merge-reports - Merge sharded reports")
add_code(doc, "npx playwright merge-reports [options] <blob-report-dir>")
add_normal(doc, "Options:")
add_code(doc, "  --config <file>       Config file path")
add_code(doc, "  --reporter <reporter> Output reporter: html | json")
add_spacer(doc)

# ── 11. Utilities ─────────────────────────────────────────────────────────────
add_heading(doc, "11. Utilities")
add_code(doc, "npx playwright uninstall          # Remove installed browsers")
add_code(doc, "npx playwright install-deps       # Install OS-level browser dependencies")
add_code(doc, "npx playwright clear-cache        # Clear Playwright build and test caches")
add_code(doc, "npx playwright init-agents        # Initialize repository agents")
add_code(doc, "npx playwright --version          # Show installed version")
add_code(doc, "npx playwright --help             # Show all commands")
add_code(doc, "npx playwright help <command>     # Help for a specific command")
add_spacer(doc)

# ── 12. Windows PowerShell Tips ───────────────────────────────────────────────
add_heading(doc, "12. Windows PowerShell Tips")
add_normal(doc, "On Windows PowerShell, set environment variables using $env: syntax:")
add_code(doc, "  # Enable Inspector while running tests")
add_code(doc, "  $env:PWDEBUG=1; npx playwright test")
add_code(doc, "")
add_code(doc, "  # Slow down actions by 1000ms")
add_code(doc, "  $env:PLAYWRIGHT_SLOW_MO=1000; npx playwright test")
add_code(doc, "")
add_code(doc, "  # Linux/Mac syntax (does NOT work on PowerShell)")
add_code(doc, "  PWDEBUG=1 npx playwright test   # WRONG on Windows")

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\ITSD\Desktop\Playwright\AI_playwright\Playwright_CLI_Reference.docx"
doc.save(output_path)
print(f"Word file saved to: {output_path}")
